import html
import io
import json
import os
import re
import urllib.error
import urllib.request
import xml.etree.ElementTree as ET
import zipfile
from collections import defaultdict
from datetime import datetime, timedelta
from pathlib import Path

import markdown
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from dotenv import load_dotenv

from prompts.format_prompt import FormatOptions
from services.diff_utils import compute_diff
from services.feedback import add_feedback, init_feedback_db
from services.formatter import FormatterService
from services.history import (
    add_history_item,
    fetch_history,
    fetch_history_item,
    init_history_db,
    remove_history_item,
    update_history_rating,
    update_history_title,
)
from services.export_settings import load_export_settings, save_export_settings
from services.presets import (
    delete_custom_preset,
    get_custom_preset,
    init_preset_db,
    list_custom_presets,
    save_custom_preset,
)
from services.ui_settings import load_ui_settings, save_ui_settings
from services.youtube_transcript import fetch_youtube_transcript


load_dotenv()
init_history_db()
init_feedback_db()
init_preset_db()

ASSETS_DIR = Path("assets")
FAVICON_PATH = ASSETS_DIR / "favicon-32.png"
SIDEBAR_LOGO_PATH = ASSETS_DIR / "logo-192.png"

st.set_page_config(
    page_title="Formatr",
    page_icon=str(FAVICON_PATH) if FAVICON_PATH.exists() else "📝",
    layout="wide",
)


THEME_STYLES = {
    "Default": {"bg": "#ffffff", "fg": "#111827", "border": "#e5e7eb"},
    "Slate": {"bg": "#0f172a", "fg": "#e2e8f0", "border": "#334155"},
    "Sepia": {"bg": "#fdf6e3", "fg": "#3f3a2f", "border": "#e6ddc6"},
}

PRESET_CONFIGS = {
    "minimal_cleanup": {
        "text_changes_level": "minimal",
        "enable_bold": True,
        "enable_italics": True,
        "enable_h1": True,
        "enable_h2": True,
        "enable_h3": True,
        "bullets_mode": "auto",
        "pull_quotes_mode": "off",
        "numbered_steps_mode": "off",
        "section_summaries_mode": "off",
        "tables_mode": "off",
        "callouts_mode": "off",
    },
    "article": {
        "text_changes_level": "thorough",
        "enable_bold": True,
        "enable_italics": True,
        "enable_h1": True,
        "enable_h2": True,
        "enable_h3": True,
        "bullets_mode": "prefer",
        "pull_quotes_mode": "prefer",
        "numbered_steps_mode": "auto",
        "section_summaries_mode": "auto",
        "tables_mode": "auto",
        "callouts_mode": "off",
    },
    "executive_brief": {
        "text_changes_level": "minimal",
        "enable_bold": True,
        "enable_italics": False,
        "enable_h1": False,
        "enable_h2": True,
        "enable_h3": True,
        "bullets_mode": "prefer",
        "pull_quotes_mode": "off",
        "numbered_steps_mode": "auto",
        "section_summaries_mode": "prefer",
        "tables_mode": "auto",
        "callouts_mode": "prefer",
    },
    "tutorial": {
        "text_changes_level": "thorough",
        "enable_bold": True,
        "enable_italics": True,
        "enable_h1": True,
        "enable_h2": True,
        "enable_h3": True,
        "bullets_mode": "prefer",
        "pull_quotes_mode": "off",
        "numbered_steps_mode": "prefer",
        "section_summaries_mode": "auto",
        "tables_mode": "auto",
        "callouts_mode": "prefer",
    },
}

FORMAT_SETTING_KEYS = [
    "text_changes_level",
    "enable_bold",
    "enable_italics",
    "enable_h1",
    "enable_h2",
    "enable_h3",
    "bullets_mode",
    "pull_quotes_mode",
    "numbered_steps_mode",
    "section_summaries_mode",
    "tables_mode",
    "callouts_mode",
]

DEFAULT_EXPORT_SETTINGS = {
    "export_use_project_template": False,
    "export_project_template_path": "templates/base_template.docx",
    "export_project_template_choice": "(Manual path)",
    "export_use_custom_heading_styles": False,
    "export_h1_font": "",
    "export_h1_size": 16,
    "export_h2_font": "",
    "export_h2_size": 14,
    "export_h3_font": "",
    "export_h3_size": 12,
}

PERSISTED_UI_SETTINGS = [
    "provider",
    "free_models_only",
    "openrouter_model",
    "openai_model",
    "gemini_model",
    "font",
    "theme",
    "history_sort_order",
    "text_changes_level",
    "enable_bold",
    "enable_italics",
    "enable_h1",
    "enable_h2",
    "enable_h3",
    "bullets_mode",
    "pull_quotes_mode",
    "numbered_steps_mode",
    "section_summaries_mode",
    "tables_mode",
    "callouts_mode",
    "format_preset",
]


def initialize_state() -> None:
    defaults = {
        "raw_text": "",
        "youtube_video_id": "",
        "formatted_text": "",
        "stream_title": "Formatted Result",
        "stream_title_edit": "Formatted Result",
        "text_changes_level": "minimal",
        "enable_bold": True,
        "enable_italics": True,
        "enable_h1": True,
        "enable_h2": True,
        "enable_h3": True,
        "bullets_mode": "auto",
        "pull_quotes_mode": "auto",
        "numbered_steps_mode": "auto",
        "section_summaries_mode": "auto",
        "tables_mode": "off",
        "callouts_mode": "off",
        "format_preset": "custom",
        "last_applied_preset": "custom",
        "custom_preset_name": "",
        "font": "Montserrat",
        "theme": "Default",
        "show_diff": False,
        "current_doc_id": None,
        "copied": False,
        "history_sort_order": "Newest first",
        "provider": os.getenv("OPENROUTER_FAMILY", "openai"),
        "free_models_only": False,
        "openrouter_model": os.getenv("OPENROUTER_MODEL", "openai/gpt-4o-mini"),
        "openai_model": os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
        "gemini_model": os.getenv("GEMINI_MODEL", "gemini-2.5-flash"),
        "export_use_custom_heading_styles": False,
        "export_h1_font": "",
        "export_h1_size": 16,
        "export_h2_font": "",
        "export_h2_size": 14,
        "export_h3_font": "",
        "export_h3_size": 12,
        "_export_settings_loaded": False,
        "_ui_settings_loaded": False,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def initialize_export_settings() -> None:
    if st.session_state.get("_export_settings_loaded"):
        return
    for key, default_value in DEFAULT_EXPORT_SETTINGS.items():
        if key not in st.session_state:
            st.session_state[key] = default_value
    saved = load_export_settings()
    for key in DEFAULT_EXPORT_SETTINGS:
        if key in saved:
            st.session_state[key] = saved[key]
    st.session_state["_export_settings_loaded"] = True


def initialize_ui_settings() -> None:
    if st.session_state.get("_ui_settings_loaded"):
        return
    saved = load_ui_settings()
    for key in PERSISTED_UI_SETTINGS:
        if key in saved and key in st.session_state:
            st.session_state[key] = saved[key]
    st.session_state["_ui_settings_loaded"] = True


def persist_ui_settings() -> None:
    payload = {key: st.session_state.get(key) for key in PERSISTED_UI_SETTINGS}
    save_ui_settings(payload)


def persist_export_settings() -> None:
    save_export_settings({key: st.session_state.get(key) for key in DEFAULT_EXPORT_SETTINGS})


def reset_export_settings_to_defaults() -> None:
    for key, value in DEFAULT_EXPORT_SETTINGS.items():
        st.session_state[key] = value
    persist_export_settings()


def get_word_count(text: str) -> int:
    return len([token for token in text.strip().split() if token]) if text.strip() else 0


PROVIDER_DISPLAY_NAMES = {
    "openai": "OpenAI",
    "google": "Google",
    "anthropic": "Anthropic",
    "meta-llama": "Meta",
    "mistralai": "Mistral",
    "qwen": "Qwen",
    "deepseek": "DeepSeek",
    "x-ai": "xAI",
    "moonshotai": "Moonshot",
    "cohere": "Cohere",
    "microsoft": "Microsoft",
    "amazon": "Amazon",
}

PREFERRED_PROVIDER_ORDER = [
    "openai",
    "google",
    "anthropic",
    "meta-llama",
    "mistralai",
    "qwen",
    "deepseek",
    "x-ai",
    "moonshotai",
    "cohere",
    "microsoft",
    "amazon",
]


def provider_label(provider_slug: str) -> str:
    return PROVIDER_DISPLAY_NAMES.get(provider_slug, provider_slug.replace("-", " ").title())


@st.cache_data(ttl=3600, show_spinner=False)
def fetch_openrouter_models(api_key_for_cache: str) -> list[str]:
    req = urllib.request.Request("https://openrouter.ai/api/v1/models")
    req.add_header("Accept", "application/json")
    if api_key_for_cache:
        req.add_header("Authorization", f"Bearer {api_key_for_cache}")
    try:
        with urllib.request.urlopen(req, timeout=20) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = ""
        try:
            detail = exc.read().decode("utf-8")
        except Exception:  # noqa: BLE001
            detail = str(exc)
        raise RuntimeError(f"OpenRouter models request failed ({exc.code}): {detail}") from exc
    except (urllib.error.URLError, TimeoutError, json.JSONDecodeError) as exc:
        raise RuntimeError(f"OpenRouter models request failed: {exc}") from exc

    model_ids: list[str] = []
    for item in payload.get("data", []):
        model_id = str(item.get("id", "")).strip()
        if model_id and "/" in model_id:
            model_ids.append(model_id)
    return sorted(set(model_ids), key=str.lower)


def get_provider_families(model_ids: list[str]) -> list[str]:
    discovered = {model_id.split("/", 1)[0] for model_id in model_ids if "/" in model_id}
    ordered = [family for family in PREFERRED_PROVIDER_ORDER if family in discovered]
    ordered.extend(sorted(discovered - set(ordered)))
    return ordered


def get_models_for_family(model_ids: list[str], family: str) -> list[str]:
    return [model_id for model_id in model_ids if model_id.startswith(f"{family}/")]


def provider_label_with_count(provider_slug: str, model_ids: list[str]) -> str:
    if provider_slug == "__all__":
        return f"All providers ({len(model_ids)})"
    return f"{provider_label(provider_slug)} ({len(get_models_for_family(model_ids, provider_slug))})"


def list_project_templates() -> list[str]:
    templates_dir = Path("templates")
    if not templates_dir.exists() or not templates_dir.is_dir():
        return []
    found: list[str] = []
    for ext in ("*.docx", "*.dotx", "*.dotm"):
        for path in templates_dir.rglob(ext):
            if path.is_file():
                found.append(str(path).replace("\\", "/"))
    return sorted(set(found))


def inspect_template_styles(template_path: str) -> dict:
    candidate = Path((template_path or "").strip())
    if not candidate.exists() or not candidate.is_file():
        return {"ok": False, "error": f"Template not found: {candidate}"}

    high_level_error = ""
    xml_error = ""
    try:
        doc = Document(str(candidate))
    except Exception as exc:
        return {"ok": False, "error": f"Template load failed: {exc}"}

    paragraph_styles = []
    all_styles = []
    heading_found = {"Heading 1": False, "Heading 2": False, "Heading 3": False}
    heading_aliases = {
        "heading 1": "Heading 1",
        "heading1": "Heading 1",
        "heading 2": "Heading 2",
        "heading2": "Heading 2",
        "heading 3": "Heading 3",
        "heading3": "Heading 3",
    }

    try:
        for style in doc.styles:
            style_id = str(getattr(style, "style_id", "") or "")
            style_name = str(getattr(style, "name", "") or "")
            style_type = getattr(style, "type", None)
            style_type_name = {
                WD_STYLE_TYPE.PARAGRAPH: "paragraph",
                WD_STYLE_TYPE.CHARACTER: "character",
                WD_STYLE_TYPE.TABLE: "table",
                WD_STYLE_TYPE.LIST: "list/numbering",
            }.get(style_type, str(style_type))

            all_styles.append(
                {"style_id": style_id, "name": style_name, "style_type": style_type_name}
            )

            if style_type == WD_STYLE_TYPE.PARAGRAPH:
                paragraph_styles.append({"style_id": style_id, "name": style_name})

            normalized_name = style_name.strip().lower().replace("-", " ")
            normalized_id = style_id.strip().lower().replace("-", " ")
            for alias, canonical in heading_aliases.items():
                if alias in {normalized_name, normalized_id}:
                    heading_found[canonical] = True
    except Exception as exc:
        high_level_error = str(exc)

    # Fallback for templates where python-docx style enumeration is empty:
    # read styles directly from word/styles.xml
    if not all_styles:
        try:
            with zipfile.ZipFile(candidate, "r") as zf:
                raw_xml = zf.read("word/styles.xml")
            root = ET.fromstring(raw_xml)
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for style_el in root.findall("w:style", ns):
                style_id = style_el.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId", "")
                style_type = style_el.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "")
                name_el = style_el.find("w:name", ns)
                style_name = ""
                if name_el is not None:
                    style_name = name_el.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "")

                style_type_name = {
                    "paragraph": "paragraph",
                    "character": "character",
                    "table": "table",
                    "numbering": "list/numbering",
                }.get(style_type, style_type or "unknown")

                row = {
                    "style_id": style_id,
                    "name": style_name,
                    "style_type": style_type_name,
                    "source": "styles.xml",
                }
                all_styles.append(row)
                if style_type == "paragraph":
                    paragraph_styles.append({"style_id": style_id, "name": style_name})

                normalized_name = style_name.strip().lower().replace("-", " ")
                normalized_id = style_id.strip().lower().replace("-", " ")
                for alias, canonical in heading_aliases.items():
                    if alias in {normalized_name, normalized_id}:
                        heading_found[canonical] = True
        except Exception as exc:
            xml_error = str(exc)

    # Mark python-docx source for rows gathered from the high-level API.
    if all_styles and "source" not in all_styles[0]:
        for row in all_styles:
            row["source"] = "python-docx"

    return {
        "ok": True,
        "path": str(candidate),
        "heading_found": heading_found,
        "paragraph_styles": paragraph_styles,
        "all_styles": all_styles,
        "high_level_count": len(all_styles),
        "paragraph_count": len(paragraph_styles),
        "high_level_error": high_level_error,
        "xml_error": xml_error,
        "file_size_bytes": candidate.stat().st_size if candidate.exists() else 0,
    }


def _style_rows_for_display(rows: list[dict], include_type: bool = False) -> list[dict]:
    display_rows: list[dict] = []
    for idx, row in enumerate(rows, start=1):
        item = {
            "#": idx,
            "style_id": (row.get("style_id") or "").strip() or "(blank)",
            "name": (row.get("name") or "").strip() or "(blank)",
        }
        if include_type:
            item["style_type"] = (row.get("style_type") or "").strip() or "(blank)"
            item["source"] = (row.get("source") or "").strip() or "(blank)"
        display_rows.append(item)
    return display_rows


def _style_rows_to_tsv(rows: list[dict]) -> str:
    if not rows:
        return "(no rows)"
    headers = list(rows[0].keys())
    lines = ["\t".join(headers)]
    for row in rows:
        lines.append("\t".join(str(row.get(h, "")) for h in headers))
    return "\n".join(lines)


def group_history(items: list[dict], sort_order: str) -> list[tuple[str, list[dict]]]:
    now = datetime.now()
    today = datetime(now.year, now.month, now.day)
    yesterday = today - timedelta(days=1)
    week_ago = today - timedelta(days=7)

    sorted_items = sorted(
        items,
        key=lambda item: item.get("created_at", ""),
        reverse=(sort_order == "Newest first"),
    )
    groups: dict[str, list[dict]] = defaultdict(list)
    for item in sorted_items:
        created = datetime.fromisoformat(item["created_at"])
        if created >= today:
            groups["Today"].append(item)
        elif created >= yesterday:
            groups["Yesterday"].append(item)
        elif created >= week_ago:
            groups["Last 7 days"].append(item)
        else:
            groups["Older"].append(item)

    order_newest = ["Today", "Yesterday", "Last 7 days", "Older"]
    order_oldest = list(reversed(order_newest))
    order = order_newest if sort_order == "Newest first" else order_oldest
    return [(label, groups[label]) for label in order if groups.get(label)]


def copy_to_clipboard(text: str) -> None:
    escaped = json.dumps(text)
    components.html(
        f"""
        <script>
        navigator.clipboard.writeText({escaped});
        </script>
        """,
        height=0,
    )


def render_diff(diff_result: dict) -> str:
    chunks = []
    for token in diff_result["tokens"]:
        value = html.escape(token["value"])
        if not value:
            continue
        if token["type"] == "insert":
            chunks.append(f'<span style="background:#d1fae5;">{value}</span>')
        elif token["type"] == "delete":
            chunks.append(
                f'<span style="background:#fee2e2;text-decoration:line-through;">{value}</span>'
            )
        else:
            chunks.append(value)
    return " ".join(chunks)


def build_export_basename(title: str) -> str:
    base = (title or "formatted_result").strip().lower()
    base = re.sub(r"[^a-z0-9]+", "_", base).strip("_")
    return base or "formatted_result"


def _strip_inline_markdown(text: str) -> str:
    cleaned = re.sub(r"`([^`]*)`", r"\1", text)
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
    cleaned = re.sub(r"_([^_]+)_", r"\1", cleaned)
    return html.unescape(cleaned)


def _parse_markdown_table_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [_strip_inline_markdown(cell.strip()) for cell in raw.split("|")]


def _is_markdown_table_separator(line: str) -> bool:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    parts = [part.strip() for part in raw.split("|")]
    if not parts:
        return False
    return all(re.fullmatch(r":?-{3,}:?", part or "") for part in parts)


def _parse_markdown_table_alignment(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    parts = [part.strip() for part in raw.split("|")]
    alignments: list[str] = []
    for part in parts:
        if part.startswith(":") and part.endswith(":"):
            alignments.append("center")
        elif part.endswith(":"):
            alignments.append("right")
        else:
            alignments.append("left")
    return alignments


def _parse_markdown_table_alignment_details(line: str) -> list[tuple[str, bool]]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    parts = [part.strip() for part in raw.split("|")]

    details: list[tuple[str, bool]] = []
    for part in parts:
        explicit = ":" in part
        if part.startswith(":") and part.endswith(":"):
            details.append(("center", True))
        elif part.endswith(":"):
            details.append(("right", True))
        elif part.startswith(":"):
            details.append(("left", True))
        else:
            details.append(("left", explicit))
    return details


def _apply_cell_alignment(cell, alignment: str) -> None:
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    target = mapping.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph in cell.paragraphs:
        paragraph.alignment = target


def _is_numeric_like(cell_text: str) -> bool:
    text = (cell_text or "").strip()
    if not text or not re.search(r"\d", text):
        return False

    # Accept common finance-like numeric tokens, for example:
    # $2.5M, 12%, 150bps, 18x, USD400k, ~60,000, <= 1.2bn
    currency_codes = r"(?:usd|eur|gbp|cad|aud|jpy|chf|inr|cny|rmb)"
    numeric_core = r"(?:\d{1,3}(?:,\d{3})+|\d+)(?:\.\d+)?"
    finance_suffix = r"(?:%|x|bp|bps|k|m|mm|mn|bn|b|t)?"
    token_pattern = re.compile(r"^(?:[<>]=?)$")
    finance_token_pattern = re.compile(
        rf"(?i)^(?:{currency_codes})?"
        rf"[<>~≈]?"
        rf"\(?[+-]?"
        rf"(?:\$|€|£|¥)?"
        rf"{numeric_core}"
        rf"{finance_suffix}"
        rf"\)?"
        rf"(?:{currency_codes})?$"
    )

    normalized = html.unescape(text)
    normalized = normalized.replace("&rarr;", "->").replace("→", "->")
    normalized = re.sub(r"(?<=\d)\s*[-–]\s*(?=\d)", " ", normalized)
    normalized = normalized.replace("=", " ").replace(":", " ")
    normalized = normalized.replace("->", " ")
    normalized = normalized.strip()

    raw_tokens = [token.strip(".,;[]{}") for token in re.split(r"\s+", normalized) if token.strip()]
    if not raw_tokens:
        return False

    numeric_tokens = 0
    relevant_tokens = 0
    for token in raw_tokens:
        if token in {"-", "/", "|"}:
            continue
        relevant_tokens += 1

        # Standalone comparators (e.g., "<= 2.5M") are allowed context tokens.
        if token_pattern.fullmatch(token):
            numeric_tokens += 1
            continue
        if finance_token_pattern.fullmatch(token):
            numeric_tokens += 1

    if relevant_tokens == 0:
        return False

    # Require most meaningful tokens to look finance-numeric.
    return numeric_tokens / relevant_tokens >= 0.7


def _detect_numeric_columns(table_rows: list[list[str]], col_count: int) -> set[int]:
    numeric_columns: set[int] = set()
    for col_idx in range(col_count):
        non_empty = 0
        numeric = 0
        for row in table_rows:
            value = row[col_idx] if col_idx < len(row) else ""
            if not value.strip():
                continue
            non_empty += 1
            if _is_numeric_like(value):
                numeric += 1
        if non_empty > 0 and numeric / non_empty >= 0.6:
            numeric_columns.add(col_idx)
    return numeric_columns


def _setext_heading_level(line: str) -> int | None:
    stripped = line.strip()
    if re.fullmatch(r"={3,}", stripped):
        return 1
    if re.fullmatch(r"-{3,}", stripped):
        return 2
    return None


def _first_markdown_heading_text(markdown_text: str) -> str:
    lines = (markdown_text or "").splitlines()
    for idx, raw_line in enumerate(lines):
        stripped = raw_line.strip()
        if not stripped:
            continue
        atx_match = re.match(r"^#{1,6}\s+(.*)$", stripped)
        if atx_match:
            return _strip_inline_markdown(atx_match.group(1)).strip().lower()
        if idx + 1 < len(lines):
            setext_level = _setext_heading_level(lines[idx + 1])
            if setext_level is not None:
                return _strip_inline_markdown(stripped).strip().lower()
        break
    return ""


def build_markdown_export_text(document_title: str, markdown_text: str) -> str:
    normalized_title = (document_title or "").strip()
    body = markdown_text or ""
    if not normalized_title:
        return body
    first_heading = _first_markdown_heading_text(body)
    if first_heading and first_heading == normalized_title.lower():
        return body
    if not body.strip():
        return f"# {normalized_title}\n"
    return f"# {normalized_title}\n\n{body}"


def _apply_heading_style(run, style_config: dict | None) -> None:
    if not style_config:
        return
    font_name = (style_config.get("font_name") or "").strip()
    font_size = style_config.get("font_size")
    if font_name:
        _force_font_name(run.font, font_name)
    if isinstance(font_size, (int, float)) and font_size > 0:
        run.font.size = Pt(float(font_size))


def _force_font_name(font_obj, font_name: str) -> None:
    if not font_name:
        return
    font_obj.name = font_name
    rpr = font_obj.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    # Remove theme bindings that can force Calibri-like defaults.
    for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "csTheme", "cstheme"):
        qname = qn(f"w:{theme_attr}")
        if qname in rfonts.attrib:
            del rfonts.attrib[qname]
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _apply_heading_styles_to_document(doc: Document, heading_styles: dict[int, dict] | None) -> None:
    if not heading_styles:
        return
    for level, cfg in heading_styles.items():
        if level not in (1, 2, 3):
            continue
        style_name = f"Heading {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        font_name = (cfg.get("font_name") or "").strip()
        font_size = cfg.get("font_size")
        if font_name:
            _force_font_name(style.font, font_name)
        if isinstance(font_size, (int, float)) and font_size > 0:
            style.font.size = Pt(float(font_size))


def _add_heading_paragraph(
    doc: Document,
    heading_text: str,
    level: int,
    heading_styles: dict[int, dict] | None = None,
) -> None:
    paragraph = None
    style_candidates = [f"Heading {level}", f"Heading{level}", f"heading {level}"]
    for style_name in style_candidates:
        try:
            paragraph = doc.add_paragraph("", style=style_name)
            break
        except KeyError:
            continue

    if paragraph is None:
        # Some templates remove/rename built-in heading styles.
        # Fall back to a regular paragraph so export never crashes.
        paragraph = doc.add_paragraph("")

    run = paragraph.add_run(heading_text)
    _apply_heading_style(run, (heading_styles or {}).get(level))

    # If no heading style exists and no explicit bold override provided,
    # ensure the heading still reads as a heading.
    if paragraph.style is None or "heading" not in (paragraph.style.name or "").lower():
        run.bold = True


def _add_paragraph_with_style_fallback(
    doc: Document,
    text: str,
    style_candidates: list[str],
    bold_fallback: bool = False,
    fallback_prefix: str = "",
) :
    paragraph = None
    for style_name in style_candidates:
        try:
            paragraph = doc.add_paragraph("", style=style_name)
            break
        except KeyError:
            continue

    if paragraph is None:
        paragraph = doc.add_paragraph("")

    rendered_text = f"{fallback_prefix}{text}" if paragraph.style is None and fallback_prefix else text
    run = paragraph.add_run(rendered_text)
    if paragraph.style is None and bold_fallback:
        run.bold = True
    return paragraph


def _compute_markdown_list_indent_level(line: str) -> int:
    expanded = (line or "").replace("\t", "    ")
    leading_spaces = len(expanded) - len(expanded.lstrip(" "))
    return max(0, leading_spaces // 2)


def _apply_list_indentation(paragraph, indent_level: int) -> None:
    if paragraph is None or indent_level <= 0:
        return
    # Keep nested list hierarchy visible even if template list styles are missing.
    paragraph.paragraph_format.left_indent = Pt(18 * indent_level)


def _apply_table_style_with_fallback(table, style_candidates: list[str]) -> None:
    for style_name in style_candidates:
        try:
            table.style = style_name
            return
        except KeyError:
            continue
    # Some templates remove built-in table styles.
    # Leave default table style instead of crashing export.


def markdown_to_docx_bytes(
    markdown_text: str,
    heading_styles: dict[int, dict] | None = None,
    document_title: str = "",
    template_path: str = "",
) -> bytes:
    doc = Document()
    normalized_template_path = (template_path or "").strip()
    if normalized_template_path:
        candidate = Path(normalized_template_path)
        if candidate.exists() and candidate.is_file():
            try:
                doc = Document(str(candidate))
            except Exception as exc:
                raise RuntimeError(
                    f"Template could not be loaded ({candidate}). "
                    "If this is a .dotm template, try saving it as .dotx/.docx first."
                ) from exc
    _apply_heading_styles_to_document(doc, heading_styles)
    normalized_title = (document_title or "").strip()
    first_heading = _first_markdown_heading_text(markdown_text)
    if normalized_title and normalized_title.lower() != first_heading:
        _add_heading_paragraph(
            doc=doc,
            heading_text=normalized_title,
            level=1,
            heading_styles=heading_styles,
        )

    lines = (markdown_text or "").splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # Setext headings:
        # Heading text
        # =======   (H1)  or  -------   (H2)
        if i + 1 < len(lines):
            setext_level = _setext_heading_level(lines[i + 1])
            if setext_level is not None and "|" not in stripped:
                heading_text = _strip_inline_markdown(stripped)
                _add_heading_paragraph(
                    doc=doc,
                    heading_text=heading_text,
                    level=setext_level,
                    heading_styles=heading_styles,
                )
                i += 2
                continue

        # Markdown table block:
        # | header | ...
        # | :--- | ---: |
        # | row1  | ...  |
        if (
            "|" in stripped
            and i + 1 < len(lines)
            and "|" in lines[i + 1]
            and _is_markdown_table_separator(lines[i + 1])
        ):
            header_cells = _parse_markdown_table_row(lines[i])
            alignment_details = _parse_markdown_table_alignment_details(lines[i + 1])
            col_count = max(1, len(header_cells))
            table_rows: list[list[str]] = []

            i += 2  # Skip header + separator
            while i < len(lines):
                row_line = lines[i].strip()
                if not row_line or "|" not in row_line:
                    break
                table_rows.append(_parse_markdown_table_row(lines[i]))
                i += 1

            alignments: list[str] = []
            explicit_flags: list[bool] = []
            for col_idx in range(col_count):
                if col_idx < len(alignment_details):
                    alignment, explicit = alignment_details[col_idx]
                else:
                    alignment, explicit = ("left", False)
                alignments.append(alignment)
                explicit_flags.append(explicit)

            numeric_columns = _detect_numeric_columns(table_rows, col_count)
            for col_idx in numeric_columns:
                if col_idx < len(explicit_flags) and not explicit_flags[col_idx]:
                    alignments[col_idx] = "right"

            table = doc.add_table(rows=1, cols=col_count)
            _apply_table_style_with_fallback(
                table,
                style_candidates=["Table Grid", "TableGrid", "Light Grid", "Light Shading"],
            )

            for col_idx, cell_text in enumerate(header_cells):
                if col_idx >= col_count:
                    break
                header_cell = table.rows[0].cells[col_idx]
                paragraph = header_cell.paragraphs[0]
                run = paragraph.add_run(cell_text)
                run.bold = True
                alignment = alignments[col_idx] if col_idx < len(alignments) else "left"
                _apply_cell_alignment(header_cell, alignment)

            for cells in table_rows:
                row = table.add_row().cells
                for col_idx in range(col_count):
                    row[col_idx].text = cells[col_idx] if col_idx < len(cells) else ""
                    alignment = alignments[col_idx] if col_idx < len(alignments) else "left"
                    _apply_cell_alignment(row[col_idx], alignment)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            heading_text = _strip_inline_markdown(heading_match.group(2))
            _add_heading_paragraph(
                doc=doc,
                heading_text=heading_text,
                level=level,
                heading_styles=heading_styles,
            )
            i += 1
            continue

        bullet_match = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if bullet_match:
            paragraph = _add_paragraph_with_style_fallback(
                doc=doc,
                text=_strip_inline_markdown(bullet_match.group(2)),
                style_candidates=["List Bullet", "List Paragraph", "list bullet"],
                bold_fallback=False,
                fallback_prefix="- ",
            )
            _apply_list_indentation(paragraph, _compute_markdown_list_indent_level(line))
            i += 1
            continue

        numbered_match = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if numbered_match:
            paragraph = _add_paragraph_with_style_fallback(
                doc=doc,
                text=_strip_inline_markdown(numbered_match.group(3)),
                style_candidates=["List Number", "List Paragraph", "list number"],
                bold_fallback=False,
                fallback_prefix=f"{numbered_match.group(2)}. ",
            )
            _apply_list_indentation(paragraph, _compute_markdown_list_indent_level(line))
            i += 1
            continue

        quote_match = re.match(r"^>\s+(.*)$", stripped)
        if quote_match:
            text = _strip_inline_markdown(quote_match.group(1))
            p = None
            for style_name in ("Intense Quote", "Quote", "quote"):
                try:
                    p = doc.add_paragraph("", style=style_name)
                    break
                except KeyError:
                    continue
            if p is None:
                p = doc.add_paragraph("")
            run = p.add_run(text)
            if p.style is None:
                run.italic = True
            i += 1
            continue

        doc.add_paragraph(_strip_inline_markdown(stripped))
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def apply_preset(preset_name: str) -> None:
    if preset_name not in PRESET_CONFIGS:
        return
    for key, value in PRESET_CONFIGS[preset_name].items():
        st.session_state[key] = value


def apply_settings(settings: dict) -> None:
    for key in FORMAT_SETTING_KEYS:
        if key in settings:
            st.session_state[key] = settings[key]


def get_current_settings() -> dict:
    return {key: st.session_state.get(key) for key in FORMAT_SETTING_KEYS}


initialize_state()
initialize_export_settings()
initialize_ui_settings()

st.title("Formatr")
st.caption("Streamlit formatter clone with OpenRouter model routing, history, diff, and feedback.")

with st.sidebar:
    if SIDEBAR_LOGO_PATH.exists():
        st.image(str(SIDEBAR_LOGO_PATH), width=108)
    st.subheader("Connection")
    api_key = st.text_input(
        "OpenRouter API key",
        value=st.session_state.get("openrouter_api_key", os.getenv("OPENROUTER_API_KEY", "")),
        type="password",
        help="Stored in session only.",
    )
    st.session_state["openrouter_api_key"] = api_key

    model_ids: list[str] = []
    models_error = ""
    try:
        model_ids = fetch_openrouter_models(api_key.strip())
    except RuntimeError as exc:
        models_error = str(exc)
    if not model_ids:
        model_ids = ["openai/gpt-4o-mini"]
        if models_error:
            st.warning(models_error)

    free_only = st.toggle(
        "Free models only",
        value=bool(st.session_state.get("free_models_only", False)),
        help="Show only models whose OpenRouter id ends with 'free'.",
    )
    st.session_state["free_models_only"] = free_only
    if free_only:
        free_models = [model_id for model_id in model_ids if model_id.lower().endswith("free")]
        if free_models:
            model_ids = free_models
        else:
            st.info("No free models found in the current model list.")

    provider_families = ["__all__", *(get_provider_families(model_ids) or ["openai"])]
    provider_default = st.session_state.get("provider", provider_families[0])
    if provider_default not in provider_families:
        provider_default = provider_families[0]

    provider = st.selectbox(
        "Provider",
        options=provider_families,
        format_func=lambda p: provider_label_with_count(p, model_ids),
        index=provider_families.index(provider_default),
    )
    st.session_state["provider"] = provider

    family_models = model_ids if provider == "__all__" else (get_models_for_family(model_ids, provider) or model_ids)
    model_default = st.session_state.get("openrouter_model", family_models[0])
    if model_default not in family_models:
        model_default = family_models[0]
    selected_model = st.selectbox(
        "Model",
        options=family_models,
        index=family_models.index(model_default),
        help="Live list from OpenRouter /models.",
    )
    st.session_state["openrouter_model"] = selected_model

main_col, history_col = st.columns([3, 2], gap="large")

with main_col:
    st.subheader("Input")
    transcript_cols = st.columns([2.5, 1], gap="small")
    st.session_state.youtube_video_id = transcript_cols[0].text_input(
        "YouTube video ID",
        value=st.session_state.youtube_video_id,
        placeholder="e.g. dQw4w9WgXcQ",
    )
    if transcript_cols[1].button("Load Transcript", width="stretch"):
        try:
            st.session_state.raw_text = fetch_youtube_transcript(st.session_state.youtube_video_id)
        except Exception as exc:  # noqa: BLE001
            st.error(str(exc))
        else:
            st.success("Transcript loaded into Input.")
    st.session_state.raw_text = st.text_area(
        "Paste your text here...",
        value=st.session_state.raw_text,
        height=220,
    )
    st.caption(f"{get_word_count(st.session_state.raw_text)} words")

    with st.expander("Formatting options", expanded=True):
        custom_preset_names = list_custom_presets()
        custom_preset_options = [f"custom:{name}" for name in custom_preset_names]
        preset_options = [
            "custom",
            "minimal_cleanup",
            "article",
            "executive_brief",
            "tutorial",
            *custom_preset_options,
        ]
        preset_labels = {
            "custom": "Custom",
            "minimal_cleanup": "Minimal Cleanup",
            "article": "Article",
            "executive_brief": "Executive Brief",
            "tutorial": "Tutorial",
        }
        if st.session_state.format_preset not in preset_options:
            st.session_state.format_preset = "custom"

        st.session_state.format_preset = st.selectbox(
            "Preset",
            options=preset_options,
            index=preset_options.index(st.session_state.format_preset),
            format_func=lambda p: (
                preset_labels[p]
                if p in preset_labels
                else f"Custom: {p.removeprefix('custom:')}"
            ),
            help="Apply a preset, then tweak any option as needed.",
        )
        if st.session_state.format_preset != st.session_state.last_applied_preset:
            if st.session_state.format_preset in PRESET_CONFIGS:
                apply_preset(st.session_state.format_preset)
            elif st.session_state.format_preset.startswith("custom:"):
                preset_name = st.session_state.format_preset.removeprefix("custom:")
                preset_settings = get_custom_preset(preset_name)
                if preset_settings:
                    apply_settings(preset_settings)
                else:
                    st.warning(f'Custom preset "{preset_name}" was not found.')
                    st.session_state.format_preset = "custom"
            st.session_state.last_applied_preset = st.session_state.format_preset

        st.markdown("**Custom preset management**")
        preset_name_col, save_col, delete_col = st.columns([2.2, 1, 1])
        st.session_state.custom_preset_name = preset_name_col.text_input(
            "Preset name",
            value=st.session_state.custom_preset_name,
            placeholder="e.g. Product Brief",
            label_visibility="collapsed",
        )
        if save_col.button("Save preset", width="stretch"):
            name = st.session_state.custom_preset_name.strip()
            if not name:
                st.warning("Enter a custom preset name first.")
            else:
                save_custom_preset(name=name, settings=get_current_settings())
                st.session_state.format_preset = f"custom:{name}"
                st.session_state.last_applied_preset = st.session_state.format_preset
                st.success(f'Saved preset "{name}".')
                st.rerun()

        selected_preset = st.session_state.format_preset
        can_delete_selected = selected_preset.startswith("custom:")
        if delete_col.button(
            "Delete preset",
            width="stretch",
            disabled=not can_delete_selected,
        ):
            name = selected_preset.removeprefix("custom:")
            delete_custom_preset(name)
            st.session_state.format_preset = "custom"
            st.session_state.last_applied_preset = "custom"
            st.success(f'Deleted preset "{name}".')
            st.rerun()

        st.session_state.text_changes_level = st.selectbox(
            "Text changes",
            options=["none", "minimal", "thorough"],
            format_func=lambda v: {
                "none": "Strict (structure only)",
                "minimal": "Minimal (spelling & punctuation)",
                "thorough": "Thorough (grammar & clarity)",
            }[v],
            index=["none", "minimal", "thorough"].index(st.session_state.text_changes_level),
        )
        opt_cols = st.columns(5)
        st.session_state.enable_bold = opt_cols[0].checkbox("Bold", value=st.session_state.enable_bold)
        st.session_state.enable_italics = opt_cols[1].checkbox(
            "Italic", value=st.session_state.enable_italics
        )
        st.session_state.enable_h1 = opt_cols[2].checkbox("H1", value=st.session_state.enable_h1)
        st.session_state.enable_h2 = opt_cols[3].checkbox("H2", value=st.session_state.enable_h2)
        st.session_state.enable_h3 = opt_cols[4].checkbox("H3", value=st.session_state.enable_h3)

        st.markdown("**Advanced structure**")
        mode_options = ["off", "auto", "prefer"]
        mode_help = "off: never use | auto: model decides | prefer: use when useful"

        adv_cols_1 = st.columns(3)
        st.session_state.bullets_mode = adv_cols_1[0].selectbox(
            "Bullets/lists",
            options=mode_options,
            index=mode_options.index(st.session_state.bullets_mode),
            help=mode_help,
        )
        st.session_state.pull_quotes_mode = adv_cols_1[1].selectbox(
            "Pull quotes",
            options=mode_options,
            index=mode_options.index(st.session_state.pull_quotes_mode),
            help=mode_help,
        )
        st.session_state.numbered_steps_mode = adv_cols_1[2].selectbox(
            "Numbered steps",
            options=mode_options,
            index=mode_options.index(st.session_state.numbered_steps_mode),
            help=mode_help,
        )

        adv_cols_2 = st.columns(3)
        st.session_state.section_summaries_mode = adv_cols_2[0].selectbox(
            "Section summaries",
            options=mode_options,
            index=mode_options.index(st.session_state.section_summaries_mode),
            help=mode_help,
        )
        st.session_state.tables_mode = adv_cols_2[1].selectbox(
            "Tables",
            options=mode_options,
            index=mode_options.index(st.session_state.tables_mode),
            help=mode_help,
        )
        st.session_state.callouts_mode = adv_cols_2[2].selectbox(
            "Callouts",
            options=mode_options,
            index=mode_options.index(st.session_state.callouts_mode),
            help=mode_help,
        )

    format_clicked = st.button(
        "Format",
        type="primary",
        disabled=not st.session_state.raw_text.strip(),
        width="content",
    )

    if format_clicked:
        selected_api_key = st.session_state.get("openrouter_api_key", "")
        selected_model = st.session_state.get("openrouter_model", "")

        if not selected_api_key:
            st.error("Enter your OpenRouter API key in the sidebar.")
        else:
            options = FormatOptions(
                enable_bold=st.session_state.enable_bold,
                enable_italics=st.session_state.enable_italics,
                enable_h1=st.session_state.enable_h1,
                enable_h2=st.session_state.enable_h2,
                enable_h3=st.session_state.enable_h3,
                bullets_mode=st.session_state.bullets_mode,
                pull_quotes_mode=st.session_state.pull_quotes_mode,
                numbered_steps_mode=st.session_state.numbered_steps_mode,
                section_summaries_mode=st.session_state.section_summaries_mode,
                tables_mode=st.session_state.tables_mode,
                callouts_mode=st.session_state.callouts_mode,
            )
            service = FormatterService(
                provider=st.session_state.provider,
                api_key=selected_api_key,
                model=selected_model,
                use_openrouter=True,
            )
            placeholder = st.empty()
            aggregated = ""
            try:
                with st.spinner("Formatting..."):
                    for chunk in service.stream_format_text(
                        text=st.session_state.raw_text,
                        text_changes_level=st.session_state.text_changes_level,
                        options=options,
                    ):
                        aggregated += chunk
                        placeholder.markdown(aggregated)

                st.session_state.formatted_text = aggregated
                st.session_state.stream_title = service.generate_title(
                    raw_text=st.session_state.raw_text,
                    formatted_text=aggregated,
                )
                st.session_state.stream_title_edit = st.session_state.stream_title
                st.session_state.current_doc_id = add_history_item(
                    title=st.session_state.stream_title,
                    input_text=st.session_state.raw_text,
                    formatted_text=aggregated,
                )
                st.success("Formatting complete.")
            except Exception as exc:  # noqa: BLE001
                st.error(str(exc))

    if st.session_state.formatted_text:
        toolbar_cols = st.columns([1.1, 1.1, 1.2, 1, 1, 1])

        with toolbar_cols[0]:
            st.session_state.font = st.selectbox(
                "Font",
                options=["Montserrat", "Roboto", "Times New Roman"],
                index=["Montserrat", "Roboto", "Times New Roman"].index(st.session_state.font),
                label_visibility="collapsed",
            )
        with toolbar_cols[1]:
            st.session_state.theme = st.selectbox(
                "Theme",
                options=list(THEME_STYLES.keys()),
                index=list(THEME_STYLES.keys()).index(st.session_state.theme),
                label_visibility="collapsed",
            )
        with toolbar_cols[2]:
            st.session_state.show_diff = st.toggle("Diff view", value=st.session_state.show_diff)
        with toolbar_cols[3]:
            if st.button("Copy output", width="stretch"):
                copy_to_clipboard(st.session_state.formatted_text)
                st.toast("Copied to clipboard")
        with toolbar_cols[4]:
            if st.button("👍", width="stretch") and st.session_state.current_doc_id:
                update_history_rating(st.session_state.current_doc_id, 1)
                st.toast("Thanks for the feedback")
        with toolbar_cols[5]:
            if st.button("👎", width="stretch") and st.session_state.current_doc_id:
                update_history_rating(st.session_state.current_doc_id, -1)
                st.toast("Thanks for the feedback")

        export_name = build_export_basename(st.session_state.stream_title)
        with st.expander("Word Export Settings", expanded=False):
            st.caption("Optional heading style overrides for DOCX exports. Saved automatically.")
            if st.button("Reset heading overrides", width="content"):
                reset_export_settings_to_defaults()
                st.success("Heading overrides reset to defaults.")
                st.rerun()

            st.session_state.export_use_project_template = st.toggle(
                "Use project template",
                value=st.session_state.get("export_use_project_template", False),
            )
            if st.session_state.export_use_project_template:
                discovered_templates = list_project_templates()
                if not discovered_templates:
                    st.info(
                        "No templates found in `templates/`. Add a `.docx`, `.dotx`, or `.dotm` file there."
                    )
                    if st.button("Create templates folder", key="create_templates_folder"):
                        Path("templates").mkdir(parents=True, exist_ok=True)
                        st.success("Created `templates/` folder.")
                        st.rerun()
                manual_label = "(Manual path)"
                template_options = [manual_label, *discovered_templates]

                current_choice = st.session_state.get("export_project_template_choice", manual_label)
                current_path = (
                    st.session_state.get("export_project_template_path", "") or ""
                ).strip().replace("\\", "/")
                if current_choice not in template_options:
                    current_choice = current_path if current_path in discovered_templates else manual_label

                st.session_state.export_project_template_choice = st.selectbox(
                    "Project template",
                    options=template_options,
                    index=template_options.index(current_choice),
                    help="Choose a template from templates/ or enter a manual path.",
                )

                if st.session_state.export_project_template_choice != manual_label:
                    st.session_state.export_project_template_path = (
                        st.session_state.export_project_template_choice
                    )
                else:
                    st.text_input(
                        "Template path",
                        key="export_project_template_path",
                        placeholder="templates/base_template.docx",
                        help="Path relative to project root or absolute path.",
                    )

                template_path = (st.session_state.export_project_template_path or "").strip()
                if template_path:
                    template_file = Path(template_path)
                    if template_file.exists() and template_file.is_file():
                        st.caption(f"Template found: {template_file}")
                    else:
                        st.warning(
                            "Template file not found. Export will fall back to default document style."
                        )

                    test_cols = st.columns([1.2, 2.8])
                    if test_cols[0].button("Test template styles", key="test_template_styles"):
                        st.session_state["_template_test_result"] = inspect_template_styles(template_path)
                        st.session_state["_template_test_for_path"] = template_path

                    template_test_result = st.session_state.get("_template_test_result")
                    template_test_path = st.session_state.get("_template_test_for_path", "")
                    if template_test_result and template_test_path == template_path:
                        if not template_test_result.get("ok", False):
                            test_cols[1].warning(template_test_result.get("error", "Template test failed."))
                        else:
                            test_cols[1].caption(
                                "Tested file: "
                                f"{template_test_result.get('path', template_path)} "
                                f"({template_test_result.get('file_size_bytes', 0)} bytes)"
                            )
                            heading_found = template_test_result["heading_found"]
                            h1 = "yes" if heading_found.get("Heading 1") else "no"
                            h2 = "yes" if heading_found.get("Heading 2") else "no"
                            h3 = "yes" if heading_found.get("Heading 3") else "no"
                            test_cols[1].caption(
                                f"Heading styles detected -> H1: {h1}, H2: {h2}, H3: {h3}. "
                                f"Paragraph styles: {template_test_result.get('paragraph_count', 0)}"
                            )
                            if template_test_result.get("high_level_error"):
                                st.warning(
                                    "python-docx style enumeration warning: "
                                    f"{template_test_result['high_level_error']}"
                                )
                            if template_test_result.get("xml_error"):
                                st.warning(
                                    "styles.xml fallback warning: "
                                    f"{template_test_result['xml_error']}"
                                )
                            paragraph_styles = template_test_result.get("paragraph_styles", [])
                            if paragraph_styles:
                                paragraph_rows = _style_rows_for_display(
                                    paragraph_styles, include_type=False
                                )
                                with st.expander("Paragraph styles found in template", expanded=False):
                                    st.table(paragraph_rows)
                                    with st.expander("Browser-safe plain text view", expanded=False):
                                        st.code(_style_rows_to_tsv(paragraph_rows), language="text")
                            else:
                                st.warning(
                                    "No paragraph styles were exposed by this template. "
                                    "Showing all discovered style definitions below."
                                )
                            all_style_rows = _style_rows_for_display(
                                template_test_result.get("all_styles", []),
                                include_type=True,
                            )
                            with st.expander("All styles found in template", expanded=False):
                                st.table(all_style_rows)
                                with st.expander("Browser-safe plain text view", expanded=False):
                                    st.code(_style_rows_to_tsv(all_style_rows), language="text")

            st.session_state.export_use_custom_heading_styles = st.toggle(
                "Use heading overrides",
                value=st.session_state.export_use_custom_heading_styles,
            )
            if st.session_state.export_use_custom_heading_styles:
                heading_cols = st.columns(3)
                with heading_cols[0]:
                    st.text_input("H1 font family", key="export_h1_font", placeholder="e.g. Arial")
                    st.number_input(
                        "H1 size (pt)",
                        min_value=8,
                        max_value=48,
                        step=1,
                        key="export_h1_size",
                    )
                with heading_cols[1]:
                    st.text_input("H2 font family", key="export_h2_font", placeholder="e.g. Arial")
                    st.number_input(
                        "H2 size (pt)",
                        min_value=8,
                        max_value=48,
                        step=1,
                        key="export_h2_size",
                    )
                with heading_cols[2]:
                    st.text_input("H3 font family", key="export_h3_font", placeholder="e.g. Arial")
                    st.number_input(
                        "H3 size (pt)",
                        min_value=8,
                        max_value=48,
                        step=1,
                        key="export_h3_size",
                    )

            persist_export_settings()

        markdown_export_text = build_markdown_export_text(
            st.session_state.stream_title,
            st.session_state.formatted_text,
        )
        markdown_data = markdown_export_text.encode("utf-8")
        heading_styles = None
        if st.session_state.export_use_custom_heading_styles:
            heading_styles = {
                1: {
                    "font_name": st.session_state.export_h1_font,
                    "font_size": st.session_state.export_h1_size,
                },
                2: {
                    "font_name": st.session_state.export_h2_font,
                    "font_size": st.session_state.export_h2_size,
                },
                3: {
                    "font_name": st.session_state.export_h3_font,
                    "font_size": st.session_state.export_h3_size,
                },
            }
        if st.session_state.export_use_project_template and st.session_state.export_use_custom_heading_styles:
            st.info("Heading overrides are ON, so H1/H2/H3 may differ from template heading styles.")

        template_path_to_use = (
            st.session_state.export_project_template_path
            if st.session_state.export_use_project_template
            else ""
        )
        try:
            docx_data = markdown_to_docx_bytes(
                st.session_state.formatted_text,
                heading_styles=heading_styles,
                document_title=st.session_state.stream_title,
                template_path=template_path_to_use,
            )
        except RuntimeError as exc:
            st.warning(f"{exc} Falling back to default export styles.")
            docx_data = markdown_to_docx_bytes(
                st.session_state.formatted_text,
                heading_styles=heading_styles,
                document_title=st.session_state.stream_title,
                template_path="",
            )
        export_cols = st.columns(2)
        export_cols[0].download_button(
            "Download Markdown (.md)",
            data=markdown_data,
            file_name=f"{export_name}.md",
            mime="text/markdown",
            width="stretch",
        )
        export_cols[1].download_button(
            "Download Word (.docx)",
            data=docx_data,
            file_name=f"{export_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            width="stretch",
        )

        if st.button("Reset session"):
            st.session_state.raw_text = ""
            st.session_state.youtube_video_id = ""
            st.session_state.formatted_text = ""
            st.session_state.stream_title = "Formatted Result"
            st.session_state.stream_title_edit = "Formatted Result"
            st.session_state.current_doc_id = None
            st.rerun()

        # Keep title near the rendered result block.
        title_cols = st.columns([6, 1.4])
        title_cols[0].text_input(
            "Document title",
            key="stream_title_edit",
            label_visibility="collapsed",
            placeholder="Enter document title",
        )
        save_title_disabled = not (st.session_state.stream_title_edit or "").strip()
        if title_cols[1].button("Save title", width="stretch", disabled=save_title_disabled):
            new_title = st.session_state.stream_title_edit.strip()
            st.session_state.stream_title = new_title
            if st.session_state.current_doc_id:
                update_history_title(st.session_state.current_doc_id, new_title)
            st.toast("Title saved")
            st.rerun()

        st.subheader(st.session_state.stream_title)

        if st.session_state.show_diff:
            diff_result = compute_diff(st.session_state.raw_text, st.session_state.formatted_text)
            stats = diff_result["stats"]
            st.caption(
                f"Changes: {stats['change_percent']}% | +{stats['words_added']} "
                f"-{stats['words_removed']} ~{stats['words_changed']}"
            )
            st.markdown(
                f"<div style='border:1px solid #e5e7eb;padding:12px;border-radius:8px;'>"
                f"{render_diff(diff_result)}</div>",
                unsafe_allow_html=True,
            )
        else:
            theme_cfg = THEME_STYLES[st.session_state.theme]
            html_body = markdown.markdown(
                st.session_state.formatted_text,
                extensions=["extra", "sane_lists"],
            )
            st.markdown(
                f"""
                <div style="
                    font-family:{st.session_state.font};
                    background:{theme_cfg['bg']};
                    color:{theme_cfg['fg']};
                    border:1px solid {theme_cfg['border']};
                    border-radius:10px;
                    padding:14px;
                ">
                    {html_body}
                </div>
                """,
                unsafe_allow_html=True,
            )

st.divider()
st.subheader("Feedback")
feedback_text = st.text_area("Suggest an improvement", key="feedback_text", height=80)
if st.button("Submit feedback"):
    if feedback_text.strip():
        add_feedback(feedback_text.strip())
        st.session_state.feedback_text = ""
        st.success("Thanks for your input!")
    else:
        st.warning("Feedback cannot be empty.")

persist_ui_settings()

with history_col:
    st.subheader("History")
    st.session_state.history_sort_order = st.radio(
        "Sort",
        options=["Newest first", "Oldest first"],
        index=0 if st.session_state.history_sort_order == "Newest first" else 1,
        horizontal=True,
        label_visibility="collapsed",
    )
    items = fetch_history()
    if not items:
        st.info("No history yet.")
    else:
        grouped = group_history(items, st.session_state.history_sort_order)
        for label, group_items in grouped:
            st.caption(label)
            for item in group_items:
                button_label = item["title"] or (item["input_text"][:60] if item["input_text"] else "Untitled")
                if st.button(button_label, key=f"load_{item['document_uuid']}", width="stretch"):
                    full_doc = fetch_history_item(item["document_uuid"])
                    if full_doc:
                        st.session_state.raw_text = full_doc["input_text"]
                        st.session_state.formatted_text = full_doc["formatted_text"]
                        st.session_state.stream_title = full_doc["title"]
                        st.session_state.stream_title_edit = full_doc["title"]
                        st.session_state.current_doc_id = full_doc["document_uuid"]
                        st.rerun()
                if st.button("Delete", key=f"del_{item['document_uuid']}"):
                    remove_history_item(item["document_uuid"])
                    if st.session_state.current_doc_id == item["document_uuid"]:
                        st.session_state.current_doc_id = None
                    st.rerun()
